from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
import re
import copy
import sys

def process_docx(input_path, output_path):
    doc = Document(input_path)

    # 处理普通段落
    for paragraph in doc.paragraphs:
        process_paragraph(paragraph)

    # 处理表格中的段落
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    process_paragraph(paragraph)

    doc.save(output_path)


def process_paragraph(paragraph):
    runs = list(paragraph.runs)
    for run in runs:
        original_text = run.text
        if not original_text:
            continue

        # 关键修复1：使用更精准的正则表达式
        parts = []
        # 匹配单词或非单词部分（包括空格）
        for match in re.finditer(r'([a-zA-Z]+)|([^a-zA-Z]+)', original_text):
            segment = match.group()
            if segment.isalpha():
                half = (len(segment) + 1) // 2
                parts.append((segment[:half], True))  # 前半加粗
                parts.append((segment[half:], False))  # 后半不加粗
            else:
                parts.append((segment, False))  # 非单词部分直接保留

        # 替换原Run为多个新Runs
        parent = run._element.getparent()
        index = parent.index(run._element)
        parent.remove(run._element)

        # 关键修复2：强制保留空格格式
        for text, is_bold in reversed(parts):
            new_run = OxmlElement('w:r')
            rPr = OxmlElement('w:rPr')

            # 复制原格式
            original_rPr = run._element.find(qn('w:rPr'))
            if original_rPr is not None:
                for child in original_rPr:
                    new_child = copy.deepcopy(child)
                    rPr.append(new_child)

            # 设置加粗
            b = rPr.find(qn('w:b'))
            if is_bold:
                if b is None:
                    b = OxmlElement('w:b')
                    rPr.append(b)
            else:
                if b is not None:
                    rPr.remove(b)

            new_run.append(rPr)
            t = OxmlElement('w:t')
            # 关键修复3：设置xml:space="preserve"防止空格合并
            t.set(qn('xml:space'), 'preserve')  # 确保空格保留
            t.text = text
            new_run.append(t)
            parent.insert(index, new_run)

import subprocess

def convert_to_pdf(doc_path, output_folder):
    command = [
        'libreoffice', '--headless', '--convert-to', 'pdf',
        doc_path, '--outdir', output_folder
    ]
    subprocess.run(command, check=True)

# 使用示例
# convert_to_pdf('input.docx', '/path/to/output')

input_word_file = sys.argv[1]
half_bold_tmp_word_file = input_word_file.replace('.docx','half_bold.docx')
output_path = ''

# 使用示例
# process_docx('DW_apb_i2c_databook.docx', 'half_bold_tmp.docx')
process_docx(input_word_file,half_bold_tmp_word_file)
convert_to_pdf(half_bold_tmp_word_file,output_path)

import os

# 删除文件
os.remove(half_bold_tmp_word_file)  # 将 'xxx' 替换为你要删除的文件路径

